"""
build_cache_preset.py
──────────────────────
Builds a hackathon demo preset for the Redis response cache: samples real
questions from the existing chunk-grounded QnA bank (filtered_chunks.json —
already has 2,725 pairs, no new generation needed), then splits the preset
in half:

  - CACHED half   — actually run through BOTH gate modes now
                    (redis_cache.gate_and_retrieve() for "concise",
                    redis_cache.retrieve_detailed() for "detailed") and
                    written into Redis under separate per-mode keys. A live
                    query against one of these in either mode returns
                    instantly, no LLM call — and the demo can show the two
                    modes' different answers side by side for the SAME
                    question.
  - UNCACHED half — left untouched. A live query against one of these always
                    takes the full retrieval + generation path, so the demo
                    can show the latency/token difference side by side.

Both gate modes are dynamic (see redis_cache.py's module docstring) — this
script does not pre-decide which underlying method serves a query; it just
runs the real gate and records whatever it actually picked.

Writes:
  data/output/hackathon_cache_preset.json  — machine-readable manifest
  data/output/hackathon_cache_demo.docx    — demo script (which question is
                                              in which bucket, which method
                                              each mode actually picked)

This is a one-time preset, not a live incremental cache — re-run this script
to regenerate it (e.g. weekly/monthly against newly-asked questions, per the
incremental-refresh idea) rather than expecting it to grow on its own.

Usage:
    python build_cache_preset.py                  # default 48 questions
    python build_cache_preset.py --count 60
"""

import argparse
import importlib.util
import json
import logging
import random
import sys
import time
from pathlib import Path

logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
log = logging.getLogger("build_cache_preset")

# Stage 1 and retrieval_layer each have their own config.py — a bare `import
# config` in both places collides in sys.modules (whichever imports first
# wins for the whole process). Load Stage 1's under an explicit local name
# first, same workaround run_tail.py already uses for this exact collision,
# so retrieval_layer's own `import config` (below) resolves to ITS config.py
# instead of silently reusing this one.
STAGE1_DIR = Path(__file__).parent
_spec = importlib.util.spec_from_file_location("stage1_config", STAGE1_DIR / "config.py")
config = importlib.util.module_from_spec(_spec)
_spec.loader.exec_module(config)

RETRIEVAL_DIR = Path(__file__).parent.parent / "retrieval_layer"
sys.path.insert(0, str(RETRIEVAL_DIR))

from ..retrieval_layer.redis_cache import gate_and_retrieve, retrieve_detailed, set_cached, _generate_answer  # noqa: E402  (retrieval_layer)
import config as rl_config                       # noqa: E402  (retrieval_layer's config.py, safe now)

FILTERED_CHUNKS_PATH = config.OUTPUT_DIR / "filtered_chunks.json"
SEED = 42   # fixed, so re-runs are reproducible unless the source QnA bank changes


def _load_qna_pool():
    with open(FILTERED_CHUNKS_PATH, encoding="utf-8") as f:
        chunks = json.load(f)
    pool = []
    for c in chunks:
        for pair in c.get("qna", []):
            q = (pair.get("q") or "").strip()
            if q:
                pool.append({"query": q, "chunk_id": c.get("chunk_id"), "source_doc": c.get("source_doc", "")})
    return pool


GATE_FNS = {
    "concise":  gate_and_retrieve,
    "detailed": lambda q: retrieve_detailed(q),
}


def build(count: int) -> None:
    pool = _load_qna_pool()
    log.info(f"Loaded {len(pool)} candidate questions from the existing QnA bank")

    rng = random.Random(SEED)
    sample = rng.sample(pool, min(count, len(pool)))
    rng.shuffle(sample)
    half = len(sample) // 2
    cached_bucket, uncached_bucket = sample[:half], sample[half:]

    log.info(f"Preset: {len(sample)} questions -> {len(cached_bucket)} cached (both gate modes) / {len(uncached_bucket)} uncached")

    # Seed the cache half — BOTH gate modes, real retrieval + generation,
    # written to Redis under separate per-mode keys (redis_cache.cache_key
    # includes mode, so concise/detailed never collide for the same question).
    for i, item in enumerate(cached_bucket, 1):
        item["modes"] = {}
        for mode, gate_fn in GATE_FNS.items():
            t0 = time.time()
            result = gate_fn(item["query"])
            method = result["_gated_method"]
            answer = _generate_answer(item["query"], result)
            latency = time.time() - t0
            set_cached(item["query"], method, answer, latency, mode=mode)
            item["modes"][mode] = {"method": method, "answer": answer, "seed_latency_s": round(latency, 2)}
        log.info(f"[{i}/{len(cached_bucket)}] cached both modes "
                 f"(concise->{item['modes']['concise']['method']}, "
                 f"detailed->{item['modes']['detailed']['method']}): {item['query'][:60]!r}")

    # Uncached half — left untouched in both modes. "detailed" always picks
    # merged_all (deterministic), but "concise" is a live confidence check —
    # genuinely undetermined until asked, so it's not pre-labeled here.
    for item in uncached_bucket:
        item["modes"] = {"detailed": {"method": "merged_all"}}

    manifest = {
        "generated_at": time.time(),
        "total": len(sample),
        "cached": cached_bucket,
        "uncached": uncached_bucket,
    }
    with open(rl_config.CACHE_PRESET_PATH, "w", encoding="utf-8") as f:
        json.dump(manifest, f, indent=2, ensure_ascii=False)
    log.info(f"Wrote manifest -> {rl_config.CACHE_PRESET_PATH}")

    _write_docx(cached_bucket, uncached_bucket)
    _write_markdown(cached_bucket, uncached_bucket)


def _write_markdown(cached_bucket, uncached_bucket) -> None:
    lines = ["# HealthRules Payer RAG — Cache Demo Script", ""]
    lines.append(
        "This preset demonstrates response-time/token savings from caching, and the "
        "two gate modes: **concise** (`gate_and_retrieve` — stays on cheap pipeline "
        "unless its own confidence is low, then escalates) and **detailed** "
        "(`retrieve_detailed` — always the fullest pool, merged_all). Ask a question "
        "from Table 1 (Cached) in either mode for an instant, pre-computed answer with "
        "no LLM call. Ask a question from Table 2 (Uncached) to see the full "
        "retrieval + generation path run live, for comparison."
    )
    lines.append("")

    lines.append("## Table 1 — Cached (instant response, both modes)")
    lines.append("")
    lines.append("| Question | Concise: method | Concise: latency (s) | Detailed: method | Detailed: latency (s) |")
    lines.append("|---|---|---|---|---|")
    for item in cached_bucket:
        c, d = item["modes"]["concise"], item["modes"]["detailed"]
        lines.append(f"| {item['query']} | {c['method']} | {c.get('seed_latency_s','')} | "
                      f"{d['method']} | {d.get('seed_latency_s','')} |")
    lines.append("")

    lines.append("## Table 2 — Uncached (full pipeline, for comparison)")
    lines.append("")
    lines.append("| Question | Detailed mode always uses |")
    lines.append("|---|---|")
    for item in uncached_bucket:
        lines.append(f"| {item['query']} | {item['modes']['detailed']['method']} |")
    lines.append("")

    lines.append("## Cached answers, in full (both modes)")
    for item in cached_bucket:
        lines.append("")
        lines.append(f"### {item['query']}")
        for mode in ("concise", "detailed"):
            m = item["modes"][mode]
            lines.append(f"\n**{mode}** — method={m['method']}, seed_latency={m.get('seed_latency_s','')}s")
            lines.append(f"\n> {m.get('answer','')}")

    with open(rl_config.CACHE_PRESET_MD, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    log.info(f"Wrote demo markdown -> {rl_config.CACHE_PRESET_MD}")


def _write_docx(cached_bucket, uncached_bucket) -> None:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading("HealthRules Payer RAG — Cache Demo Script", level=1)

    intro = doc.add_paragraph()
    intro.add_run(
        "This preset demonstrates response-time/token savings from caching, and the "
        "two gate modes: 'concise' (gate_and_retrieve — stays on cheap pipeline unless "
        "its own confidence is low, then escalates) and 'detailed' (retrieve_detailed — "
        "always the fullest pool, merged_all). Ask a question from Table 1 (Cached) in "
        "either mode for an instant, pre-computed answer with no LLM call. Ask a "
        "question from Table 2 (Uncached) to see the full retrieval + generation path "
        "run live, for comparison."
    )

    doc.add_heading("Table 1 — Cached (instant response, both modes)", level=2)
    t1 = doc.add_table(rows=1, cols=5)
    t1.style = "Light Grid Accent 1"
    hdr = t1.rows[0].cells
    hdr[0].text = "Question"
    hdr[1].text = "Concise: method"
    hdr[2].text = "Concise: latency (s)"
    hdr[3].text = "Detailed: method"
    hdr[4].text = "Detailed: latency (s)"
    for item in cached_bucket:
        row = t1.add_row().cells
        row[0].text = item["query"]
        row[1].text = item["modes"]["concise"]["method"]
        row[2].text = str(item["modes"]["concise"].get("seed_latency_s", ""))
        row[3].text = item["modes"]["detailed"]["method"]
        row[4].text = str(item["modes"]["detailed"].get("seed_latency_s", ""))

    doc.add_heading("Table 2 — Uncached (full pipeline, for comparison)", level=2)
    t2 = doc.add_table(rows=1, cols=2)
    t2.style = "Light Grid Accent 1"
    hdr = t2.rows[0].cells
    hdr[0].text, hdr[1].text = "Question", "Detailed mode always uses"
    for item in uncached_bucket:
        row = t2.add_row().cells
        row[0].text = item["query"]
        row[1].text = item["modes"]["detailed"]["method"]

    for style in doc.styles:
        if style.name == "Normal":
            style.font.size = Pt(10)

    doc.save(rl_config.CACHE_PRESET_DOC)
    log.info(f"Wrote demo doc -> {rl_config.CACHE_PRESET_DOC}")


def main():
    ap = argparse.ArgumentParser(description="Build the hackathon Redis cache demo preset")
    ap.add_argument("--count", type=int, default=48, help="Total questions in the preset (split half cached/uncached)")
    args = ap.parse_args()
    build(args.count)


if __name__ == "__main__":
    main()
